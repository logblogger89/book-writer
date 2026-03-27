"""Export router — generates DOCX and PDF downloads of a completed novel."""
import io
import re
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.models.db_models import Project
from app.services.artifact_service import get_all_active_artifacts

router = APIRouter(prefix="/api/projects/{project_id}/export", tags=["export"])


def _safe_filename(title: str) -> str:
    return re.sub(r"[^\w\-]", "_", title).strip("_") or "novel"


def _gather_chapters(artifacts: dict) -> list[dict]:
    """Return sorted list of {number, title, prose} dicts from edited_chapter_N artifacts."""
    chapters = []
    for key, art in artifacts.items():
        if not key.startswith("edited_chapter_"):
            continue
        suffix = key[len("edited_chapter_"):]
        if not suffix.isdigit():
            continue
        content = art.content_json if hasattr(art, "content_json") else art.get("content_json", art)
        prose = str(content.get("edited_prose") or content.get("prose_text") or "")
        title = str(content.get("title") or f"Chapter {suffix}")
        chapters.append({"number": int(suffix), "title": title, "prose": prose})
    return sorted(chapters, key=lambda c: c["number"])


@router.get("/docx")
async def export_docx(project_id: str, db: AsyncSession = Depends(get_db)):
    project = await db.get(Project, project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    artifacts = await get_all_active_artifacts(db, project_id)
    chapters = _gather_chapters(artifacts)
    if not chapters:
        raise HTTPException(404, "No completed chapters to export")

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project.title)
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()  # spacer

    # Logline if available
    if "logline" in artifacts:
        logline_art = artifacts["logline"]
        content = logline_art.content_json if hasattr(logline_art, "content_json") else logline_art.get("content_json", {})
        logline_text = str(content.get("logline") or "")
        if logline_text:
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = lp.add_run(f'"{logline_text}"')
            lr.italic = True
            lr.font.size = Pt(12)

    doc.add_page_break()

    # Chapters
    for ch in chapters:
        # Chapter heading
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"Chapter {ch['number']}: {ch['title']}")
        run.font.size = Pt(16)

        doc.add_paragraph()  # spacer after heading

        # Prose paragraphs
        for para_text in ch["prose"].split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph(para_text)
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = _safe_filename(project.title)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
    )


@router.get("/pdf")
async def export_pdf(project_id: str, db: AsyncSession = Depends(get_db)):
    project = await db.get(Project, project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    artifacts = await get_all_active_artifacts(db, project_id)
    chapters = _gather_chapters(artifacts)
    if not chapters:
        raise HTTPException(404, "No completed chapters to export")

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=3.5 * cm,
        rightMargin=3.5 * cm,
        topMargin=3 * cm,
        bottomMargin=3 * cm,
        title=project.title,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "NovaTitle", parent=styles["Title"],
        fontSize=28, leading=34, alignment=TA_CENTER, spaceAfter=20,
    )
    logline_style = ParagraphStyle(
        "NovaLogline", parent=styles["Italic"],
        fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=12,
        textColor=colors.HexColor("#444444"),
    )
    chapter_heading_style = ParagraphStyle(
        "NovaChapter", parent=styles["Heading1"],
        fontSize=16, leading=20, alignment=TA_CENTER, spaceBefore=0, spaceAfter=24,
        textColor=colors.HexColor("#1a1a2e"),
    )
    body_style = ParagraphStyle(
        "NovaBody", parent=styles["Normal"],
        fontSize=11, leading=16, alignment=TA_JUSTIFY,
        firstLineIndent=1 * cm, spaceAfter=8,
    )

    story = []

    # Title page
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph(project.title, title_style))
    story.append(Spacer(1, 1 * cm))

    if "logline" in artifacts:
        logline_art = artifacts["logline"]
        content = logline_art.content_json if hasattr(logline_art, "content_json") else logline_art.get("content_json", {})
        logline_text = str(content.get("logline") or "")
        if logline_text:
            story.append(Paragraph(f'&quot;{logline_text}&quot;', logline_style))

    story.append(PageBreak())

    # Chapters
    for ch in chapters:
        story.append(Paragraph(f"Chapter {ch['number']}: {ch['title']}", chapter_heading_style))
        story.append(Spacer(1, 0.3 * cm))

        for para_text in ch["prose"].split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            # Escape XML special chars
            safe = para_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body_style))

        story.append(PageBreak())

    doc_pdf.build(story)
    buf.seek(0)

    filename = _safe_filename(project.title)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
    )
